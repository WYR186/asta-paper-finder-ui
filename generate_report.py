#!/usr/bin/env python3
"""Generate Word document report for ASTA Paper Finder project."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ─── Styles ───────────────────────────────────────────────────────────────────
styles = doc.styles

def set_heading(paragraph, level, text):
    paragraph.style = f'Heading {level}'
    paragraph.text = text

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_paragraph(doc, text='', bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Light gray shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def add_table_row(table, cells):
    row = table.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = str(val)
    return row

# ─── TITLE PAGE ───────────────────────────────────────────────────────────────
title = doc.add_heading('ASTA Paper Finder', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Comprehensive Technical Report')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(16)
subtitle.runs[0].bold = True

date_p = doc.add_paragraph(f'Generated: {datetime.date.today().strftime("%B %d, %Y")}')
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ─── SECTION 1: PROJECT OVERVIEW ──────────────────────────────────────────────
add_heading(doc, '1. Project Overview and Purpose', 1)

doc.add_paragraph(
    'ASTA Paper Finder is a standalone, AI-driven academic paper discovery agent '
    'implemented as a REST API service. It is a derivative work (fork) of the original '
    'allenai/asta-paper-finder repository, with significant enhancements for local '
    'deployment and interactive use.'
)

doc.add_paragraph(
    'Core Purpose: The system assists users in locating academic papers based on natural '
    'language queries. Given a textual description of desired papers, it automatically '
    'decomposes the query, routes it to appropriate search strategies, executes multi-stage '
    'retrieval and relevance judgment, and returns a ranked list of relevant papers with '
    'explanatory text.'
)

add_heading(doc, 'Key Distinctions from Original', 2)
doc.add_paragraph('• The original online agent at paperfinder.allen.ai supports multi-turn interaction, progress updates, graphical widgets, and access to proprietary datasets.')
doc.add_paragraph('• This local version focuses on single-turn paper finding with stable, reproducible behavior.')
doc.add_paragraph('• This derivative adds: SQLite-based session history, local web UI, local deployment scripts, and bookmark/note management.')
doc.add_paragraph('Note: The repository is frozen and no longer maintained.')

# ─── SECTION 2: TECH STACK ────────────────────────────────────────────────────
add_heading(doc, '2. Tech Stack and Dependencies', 1)

add_heading(doc, 'Language & Runtime', 2)
doc.add_paragraph('• Python 3.12.8+ (required; tested on 3.13.9)')
doc.add_paragraph('• Package manager: uv (universal Python package manager)')

add_heading(doc, 'Web Framework & Server', 2)
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
tbl.rows[0].cells[0].text = 'Package'
tbl.rows[0].cells[1].text = 'Version / Purpose'
rows = [
    ('FastAPI', '~0.115.6 — REST API framework'),
    ('Uvicorn', '~0.34.0 — ASGI server'),
    ('Gunicorn', '~23.0.0 — WSGI server (used with UvicornWorker)'),
    ('Starlette', '~0.41.0 — Underlying ASGI layer'),
]
for r in rows:
    add_table_row(tbl, r)
doc.add_paragraph('')

add_heading(doc, 'LLM Integration', 2)
tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = 'Table Grid'
tbl2.rows[0].cells[0].text = 'Package'
tbl2.rows[0].cells[1].text = 'Purpose'
rows2 = [
    ('LangChain-Core ~1.0', 'LLM abstraction and callbacks'),
    ('LangChain-OpenAI ~1.0', 'OpenAI API integration'),
    ('google-genai >=1.10.0', 'Google Gemini API integration'),
    ('Cohere ~5.13.4', 'Reranking API client'),
    ('Tenacity ~9.1.2', 'Retry logic with exponential backoff'),
]
for r in rows2:
    add_table_row(tbl2, r)
doc.add_paragraph('')

add_heading(doc, 'Key External APIs Used', 2)
doc.add_paragraph('1. Semantic Scholar API (S2) — Primary source for paper metadata, citations, author data')
doc.add_paragraph('2. OpenAI API — GPT-4o, GPT-4 Turbo, GPT-5 Mini (with reasoning extensions)')
doc.add_paragraph('3. Google Gemini API — Gemini 2.0 Flash, Gemini 3.0 Flash (reasoning)')
doc.add_paragraph('4. Cohere API — Rerank English v3.0 for relevance scoring')
doc.add_paragraph('5. Vespa (AllenAI internal) — Dense vector retrieval (not publicly available)')

add_heading(doc, 'In-House Libraries (ai2i-* prefix)', 2)
doc.add_paragraph('• ai2i-config — TOML-based configuration with secrets loading')
doc.add_paragraph('• ai2i-di — Async-aware dependency injection framework with scope management')
doc.add_paragraph('• ai2i-chain — LLM abstraction layer with model registry')
doc.add_paragraph('• ai2i-dcollection — Document collection abstraction with lazy field loading')
doc.add_paragraph('• ai2i-common — Shared utilities')

# ─── SECTION 3: DIRECTORY STRUCTURE ──────────────────────────────────────────
add_heading(doc, '3. Directory / Folder Structure', 1)

doc.add_paragraph(
    'The project is organized as a Python monorepo managed by uv. The top-level directories are:'
)

tbl3 = doc.add_table(rows=1, cols=2)
tbl3.style = 'Table Grid'
tbl3.rows[0].cells[0].text = 'Directory'
tbl3.rows[0].cells[1].text = 'Purpose'
dirs = [
    ('agents/', 'All agent implementations (the main FastAPI application lives here)'),
    ('agents/mabool/api/', 'Root of the FastAPI service package (mabool)'),
    ('agents/mabool/api/mabool/api/', 'FastAPI route handlers'),
    ('agents/mabool/api/mabool/agents/', '14 specialized search agents + common utilities'),
    ('agents/mabool/api/mabool/data_model/', 'Pydantic/dataclass type definitions'),
    ('agents/mabool/api/mabool/services/', 'Dependency injection wiring'),
    ('agents/mabool/api/mabool/infra/', 'Infrastructure: Operative base class, StateManager'),
    ('agents/mabool/api/mabool/external_api/', 'External service clients (Cohere reranker)'),
    ('agents/mabool/api/mabool/utils/', 'Utility modules (LLM lookup, caching, logging)'),
    ('agents/mabool/api/mabool/dal/', 'Data access layer — S2 API configuration DI'),
    ('agents/mabool/api/conf/', 'TOML configuration files and .env.secret'),
    ('agents/mabool/api/static/', 'NEW: Single-file HTML/CSS/JS web UI (1681 lines)'),
    ('libs/', 'Six reusable Python packages (common, di, config, chain, dcollection)'),
    ('dev/', 'Developer tooling: Makefile utilities, custom flake8 plugins'),
]
for r in dirs:
    add_table_row(tbl3, r)
doc.add_paragraph('')

# ─── SECTION 4: ARCHITECTURE OVERVIEW ────────────────────────────────────────
add_heading(doc, '4. Architecture Overview', 1)

doc.add_paragraph(
    'The application follows a layered, async-first architecture. A user query travels '
    'through the following high-level pipeline:'
)

add_heading(doc, 'High-Level Pipeline', 2)
add_code_block(doc,
"""User Query (Natural Language)
    ↓
[FastAPI] POST /api/2/rounds  ←  RoundRequest (Pydantic validation)
    ↓
[Priority Semaphore]  max 3 concurrent searches
    ↓
[PaperFinderAgent]  Main orchestrator
    ├─ Query Analyzer (LLM) → QueryAnalysisResult
    │     intent, keywords, authors, venues, time range, field of study
    │
    ├─ Router (based on QueryType + operation_mode)
    │     BROAD_BY_DESCRIPTION / fast  → FastBroadSearchAgent
    │     BROAD_BY_DESCRIPTION / diligent → BroadSearchAgent (multi-round)
    │     SPECIFIC_BY_TITLE            → SpecificPaperByTitleAgent
    │     SPECIFIC_BY_NAME             → SpecificPaperByNameAgent
    │     BY_AUTHOR                    → SearchByAuthorsAgent
    │     METADATA_ONLY                → MetadataOnlySearchAgent
    │     REFUSAL                      → QueryRefusalAgent
    │
    ├─ Candidate Merge (deduplicate across all agents)
    │
    ├─ Relevance Judgment
    │     LLM batch scoring (gpt5mini-minimal-reasoning)
    │     75 concurrent requests · 250 papers max quota
    │
    ├─ Cohere Reranking (optional — skipped if no API key)
    │
    ├─ Final Sorting
    │     Multi-criteria: relevance + recency + centrality
    │
    └─ Response Generation
          LLM summary of top-10 papers
    ↓
JSON Response  +  NEW: saved to SQLite search_history.db
    ↓
Browser / API Client""")

add_heading(doc, 'Layer Descriptions', 2)
tbl4 = doc.add_table(rows=1, cols=2)
tbl4.style = 'Table Grid'
tbl4.rows[0].cells[0].text = 'Layer'
tbl4.rows[0].cells[1].text = 'Responsibility'
layers = [
    ('HTTP Layer (FastAPI)', 'Entry-point routing, request validation, response serialization'),
    ('Agent Layer (Operative)', 'Async state machine; composable sub-agents with typed I/O'),
    ('Search Agents (14)', 'Specialized retrieval strategies for different query intents'),
    ('LLM Abstraction (ai2i.chain)', 'Unified interface for OpenAI, Google Gemini; token tracking'),
    ('Document Collection (ai2i.dcollection)', 'Lazy-load paper fields; compositional operations'),
    ('Configuration (ai2i.config)', 'TOML + secrets; context-var per-request isolation'),
    ('DI Framework (ai2i.di)', 'Singleton/round-scoped providers; automatic dependency resolution'),
    ('Persistence (SQLite + file cache)', 'Search history, bookmarks, notes; query result cache'),
]
for r in layers:
    add_table_row(tbl4, r)
doc.add_paragraph('')

# ─── SECTION 5: MAJOR FILES ───────────────────────────────────────────────────
add_heading(doc, '5. Every Major File and Its Purpose', 1)

files = [
    ('API & Application', [
        ('mabool/api/app.py (111 lines)', 'FastAPI app factory. Sets up CORS, includes all route blueprints (rounds_v2_routes, stream_routes, history_routes), mounts static UI at /ui, registers error handlers, exposes /health → 204.'),
        ('mabool/api/round_v2_routes.py (139 lines)', 'POST /api/2/rounds handler. Wraps execution in FileBasedCache (TTL=600 s). Implements MaboolCallbackHandler (LangChain) to track token usage per model. Enforces priority semaphore (max 3 concurrent).'),
        ('mabool/api/history_routes.py (371 lines, NEW)', 'Full SQLite-backed CRUD for sessions, searches, bookmarks, paper_notes, and stats. Atomic INSERT with EXISTS guard prevents TOCTOU races. Provides /api/stats aggregated token usage.'),
        ('mabool/api/route_utils.py', 'Serialization helpers: converts DocumentCollection to JSON dict, appends token_breakdown_by_model and session_id to every response.'),
        ('mabool/api/stream_routes.py', 'WebSocket / streaming endpoint stubs (used by the original multi-turn online version).'),
    ]),
    ('Data Models', [
        ('mabool/data_model/rounds.py (41 lines)', 'RoundRequest (user query payload) and RoundContext (message_id, thread_id metadata).'),
        ('mabool/data_model/agent.py (260+ lines)', 'All typed I/O structures: QueryAnalysisSuccess/PartialSuccess/Refusal/Failure, AnalyzedQuery, ExtractedFields, QueryType enum, FieldOfStudy (25+ disciplines), AgentOutput, ExplainedAgentOutput, AggregatedMetrics.'),
        ('mabool/data_model/config.py (266 lines)', 'Frozen dataclass schema mirroring config.toml. Used by the DI framework to type-safely inject config values at runtime.'),
        ('mabool/data_model/specifications.py', 'Structured query specs: AuthorSpec, VenueSpec, TimeRangeSpec.'),
        ('mabool/data_model/ids.py', 'Typed ID wrappers: CorpusId (Semantic Scholar corpus ID).'),
    ]),
    ('Main Orchestrator Agent', [
        ('mabool/agents/paper_finder/paper_finder_agent.py (559 lines)', 'PaperFinderAgent. Routes based on QueryType + operation_mode. Aggregates results from sub-agents. Invokes relevance judgment → sorting → explanation. Handles refusals and clarifications.'),
    ]),
    ('Specialized Search Agents', [
        ('agents/query_analyzer/', 'decompose_and_analyze_query_restricted(): LLM-based parsing. Extracts intent, keywords, authors, venues, time range, field of study. Returns typed QueryAnalysisResult.'),
        ('agents/broad_search_by_keyword/', 'LLM generates search keywords. Executes S2 API searches in parallel (concurrency=10). Returns top-N papers per results_limit.'),
        ('agents/complex_search/fast_broad_search.py', 'FastBroadSearchAgent. Runs DenseAgent + SnowballAgent + BroadSearchByKeywordAgent in parallel. ~30 s execution.'),
        ('agents/complex_search/broad_search.py', 'BroadSearchAgent (diligent). Up to 3 iterative rounds; LLM refines queries from prior results. ~3 min execution.'),
        ('agents/dense/dense_agent.py', 'Vector semantic search against AllenAI Vespa. Silently skipped in public deployments (no public endpoint).'),
        ('agents/specific_paper_by_title/', 'Fuzzy title matching via LLM (Google Gemini). Returns S2 corpus IDs for matched papers.'),
        ('agents/specific_paper_by_name/', 'Author-cited reference matching using string similarity.'),
        ('agents/search_by_authors/', 'LLM-based author disambiguation + S2 author search API. Returns papers by specified authors.'),
        ('agents/metadata_only/', 'Pure metadata filtering (year, venue, field of study). MetadataPlannerAgent decomposes query into filter steps; no content retrieval.'),
        ('agents/snowball/snowball_agent.py', 'Citation graph traversal. Forward: papers citing the anchor (up to 200). Backward: papers cited by the anchor (up to 200). Parallel fetching.'),
        ('agents/llm_suggestion/', 'Direct LLM paper suggestions when S2 search is insufficient.'),
        ('agents/query_refusal/', 'Returns structured refusal for non-paper-finding intents (e.g., "search the web", "find an author\'s profile").'),
    ]),
    ('Common Agent Utilities', [
        ('agents/common/relevance_judgement_utils.py', 'Batch LLM scoring. Dynamic batching: starts at 20, grows by factor 2. 75 parallel requests. Max 250 papers quota.'),
        ('agents/common/sorting.py', 'SortPreferences + sorted_docs_by_preferences(): multi-criteria ranking combining relevance, recency, centrality, citation count.'),
        ('agents/common/explain.py', 'generate_response_summary(): LLM generates natural-language summary of top-10 papers.'),
        ('agents/common/computed_fields/', 'relevance.py, relevant_snippets.py — compute and cache derived document fields.'),
    ]),
    ('Configuration Files', [
        ('conf/config.toml (96 lines)', 'All default values for every agent (quotas, model names, concurrency, timeouts).'),
        ('conf/config.extra.fast_mode.toml', 'Overrides for fast mode (lower quotas, fewer iterations).'),
        ('conf/.env.secret', 'API keys: S2_API_KEY (required), OPENAI_API_KEY (required), COHERE_API_KEY (optional), GOOGLE_API_KEY (required).'),
    ]),
    ('Frontend', [
        ('static/index.html (1681 lines, NEW)', 'Single-file HTML5/CSS/JavaScript UI. Zero external dependencies. Implements: search form, result cards, session sidebar, bookmark management, notes, token usage stats, copy-to-clipboard, keyboard shortcuts.'),
    ]),
    ('Startup Scripts', [
        ('start.sh (NEW)', 'Production startup: runs Gunicorn + UvicornWorker on port 8000.'),
        ('dev.sh', 'Development startup with --reload flag.'),
    ]),
]

for section_title, file_list in files:
    add_heading(doc, section_title, 2)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.rows[0].cells[0].text = 'File / Module'
    tbl.rows[0].cells[1].text = 'Purpose'
    for f, desc in file_list:
        row = tbl.add_row()
        row.cells[0].text = f
        row.cells[1].text = desc
    doc.add_paragraph('')

# ─── SECTION 6: DATA FLOW ─────────────────────────────────────────────────────
add_heading(doc, '6. How Data Flows Through the System', 1)

add_heading(doc, 'Complete Request Lifecycle', 2)
add_code_block(doc,
"""1. Client sends:
   POST /api/2/rounds
   { "paper_description": "transformer attention for NLP",
     "operation_mode": "fast",
     "inserted_before": "2023-12-31" }

2. FastAPI → run_round_with_cache()
   ├── Check FileBasedCache (key = SHA256 of query+mode+anchors)
   │   Hit  → return cached JSON instantly
   │   Miss → execute pipeline:
   │
   │   a) Build PaperFinderInput
   │   b) Generate conversation_thread_id (UUID)
   │   c) Acquire round_semaphore (max 3 concurrent)
   │   d) Attach MaboolCallbackHandler (token tracking)
   │   e) PaperFinderAgent.handle_operation()
   │
   │      Phase 1 – Query Analysis
   │        LLM parses: keywords, intent, authors, time range, fields
   │        → QueryAnalysisSuccess
   │
   │      Phase 2 – Routing + Search
   │        BROAD + fast → FastBroadSearchAgent
   │          ├── DenseAgent (Vespa; skipped if unavailable)
   │          ├── SnowballAgent (citation graph)
   │          └── BroadSearchByKeywordAgent
   │               LLM generates keyword sets
   │               → parallel S2 API searches (concurrency=10)
   │               → ~100-300 candidate papers
   │
   │      Phase 3 – Candidate Merge
   │        Deduplicate by corpus_id → union collection
   │
   │      Phase 4 – Relevance Judgment
   │        LLM scores each abstract [0,1]
   │        75 parallel requests · dynamic batching (20→40→80…)
   │        Max 250 papers evaluated
   │
   │      Phase 5 – Cohere Reranking (if key present)
   │        rerank-english-v3.0 re-orders papers
   │
   │      Phase 6 – Final Sorting
   │        final_agent_score = weighted(relevance + recency + centrality)
   │        Sort descending
   │
   │      Phase 7 – Explanation
   │        LLM summarizes top-10 → response_text string
   │
   │   f) Store result in FileBasedCache (TTL=600 s)

3. route_utils.py serializes → JSONResponse

4. (NEW) Browser POSTs to /api/sessions/{id}/searches → saved in SQLite

5. Browser renders paper cards, updates session history

6. (Optional) User bookmarks paper → POST /api/bookmarks → SQLite""")

add_heading(doc, 'Data Transformation Summary', 2)
tbl5 = doc.add_table(rows=1, cols=3)
tbl5.style = 'Table Grid'
tbl5.rows[0].cells[0].text = 'Stage'
tbl5.rows[0].cells[1].text = 'Input Type'
tbl5.rows[0].cells[2].text = 'Output Type'
transforms = [
    ('HTTP Validation', 'JSON body', 'RoundRequest (Pydantic)'),
    ('Query Analysis', 'RoundRequest', 'QueryAnalysisResult (LLM-parsed)'),
    ('Search', 'QueryAnalysisResult', 'DocumentCollection (S2 papers)'),
    ('Merge', 'List[DocumentCollection]', 'DocumentCollection (deduplicated)'),
    ('Judgment', 'DocumentCollection', 'DocumentCollection + relevance scores'),
    ('Reranking', 'DocumentCollection', 'DocumentCollection (Cohere-reranked)'),
    ('Sorting', 'DocumentCollection', 'DocumentCollection (sorted by score)'),
    ('Response', 'DocumentCollection', 'ExplainedAgentOutput + LLM text'),
    ('Serialization', 'ExplainedAgentOutput', 'JSONResponse'),
    ('Persistence', 'JSONResponse', 'SQLite row (search_history.db)'),
]
for r in transforms:
    add_table_row(tbl5, r)
doc.add_paragraph('')

# ─── SECTION 7: API ENDPOINTS ─────────────────────────────────────────────────
add_heading(doc, '7. API Endpoints', 1)

endpoints = [
    ('Search', 'POST /api/2/rounds', 'Main search. Body: paper_description (required), operation_mode ("fast"/"diligent"/"infer"), anchor_corpus_ids (optional), inserted_before (optional date), read_results_from_cache (bool). Returns: doc_collection, response_text, analyzed_query, metrics, token_breakdown_by_model, session_id.'),
    ('Sessions', 'POST /api/sessions', 'Create a new conversation session. Body: {name}.'),
    ('Sessions', 'GET /api/sessions', 'List all sessions with search counts.'),
    ('Sessions', 'PATCH /api/sessions/{id}', 'Rename a session.'),
    ('Sessions', 'DELETE /api/sessions/{id}', 'Delete session + cascades to searches.'),
    ('Searches', 'POST /api/sessions/{id}/searches', 'Save a search result to a session.'),
    ('Searches', 'GET /api/sessions/{id}/searches', 'Retrieve all searches in a session.'),
    ('Searches', 'DELETE /api/sessions/{id}/searches/{sid}', 'Delete a single saved search.'),
    ('Bookmarks', 'POST /api/bookmarks', 'Bookmark a paper (corpus_id, title, authors, year, venue, url, abstract).'),
    ('Bookmarks', 'GET /api/bookmarks', 'List all bookmarked papers.'),
    ('Bookmarks', 'GET /api/bookmarks/ids', 'Return just the corpus IDs of bookmarks.'),
    ('Bookmarks', 'DELETE /api/bookmarks/{corpus_id}', 'Remove a bookmark.'),
    ('Bookmarks', 'PATCH /api/bookmarks/{corpus_id}', 'Update tags or note on a bookmark.'),
    ('Notes', 'PUT /api/notes/{corpus_id}', 'Create or update a note for a paper.'),
    ('Notes', 'GET /api/notes', 'Retrieve all notes as {corpus_id: note} map.'),
    ('Stats', 'GET /api/stats', 'Aggregated token usage by model + grand total.'),
    ('Health', 'GET /health', 'Returns 204; used by load balancers / monitoring.'),
    ('Docs', 'GET /docs', 'Swagger UI (FastAPI auto-generated).'),
    ('Web UI', 'GET /ui/', 'Custom single-file HTML/CSS/JS interface.'),
]

tbl6 = doc.add_table(rows=1, cols=3)
tbl6.style = 'Table Grid'
tbl6.rows[0].cells[0].text = 'Category'
tbl6.rows[0].cells[1].text = 'Endpoint'
tbl6.rows[0].cells[2].text = 'Description'
for r in endpoints:
    add_table_row(tbl6, r)
doc.add_paragraph('')

# ─── SECTION 8: DATABASE SCHEMA ───────────────────────────────────────────────
add_heading(doc, '8. Database Schema (SQLite — search_history.db)', 1)

add_heading(doc, 'Table: sessions', 2)
add_code_block(doc,
"""CREATE TABLE sessions (
    id         INTEGER PRIMARY KEY AUTOINCREMENT,
    created_at TEXT    NOT NULL DEFAULT (datetime('now')),
    name       TEXT    NOT NULL
);""")
doc.add_paragraph('Stores named conversation sessions. One row per session.')

add_heading(doc, 'Table: searches', 2)
add_code_block(doc,
"""CREATE TABLE searches (
    id            INTEGER PRIMARY KEY AUTOINCREMENT,
    session_id    INTEGER NOT NULL REFERENCES sessions(id) ON DELETE CASCADE,
    created_at    TEXT    NOT NULL DEFAULT (datetime('now')),
    query         TEXT    NOT NULL,
    mode          TEXT,
    before_date   TEXT,
    anchor_ids    TEXT,
    s2_session_id TEXT,
    result_count  INTEGER,
    result_json   TEXT        -- full API response JSON, can be 100+ KB
);""")
doc.add_paragraph('One row per executed search. result_json stores the complete API response. Cascades on session delete.')

add_heading(doc, 'Table: bookmarks', 2)
add_code_block(doc,
"""CREATE TABLE bookmarks (
    corpus_id  TEXT PRIMARY KEY,
    created_at TEXT    NOT NULL DEFAULT (datetime('now')),
    title      TEXT,
    authors    TEXT,           -- JSON array string
    year       INTEGER,
    venue      TEXT,
    url        TEXT,
    abstract   TEXT,
    tags       TEXT DEFAULT '[]',   -- JSON array string
    note       TEXT DEFAULT ''
);""")
doc.add_paragraph('User-saved papers. corpus_id is the Semantic Scholar corpus ID (globally unique).')

add_heading(doc, 'Table: paper_notes', 2)
add_code_block(doc,
"""CREATE TABLE paper_notes (
    corpus_id  TEXT PRIMARY KEY,
    note       TEXT NOT NULL DEFAULT '',
    updated_at TEXT NOT NULL DEFAULT (datetime('now'))
);""")
doc.add_paragraph('Separate per-paper notes table. Upsert operation on write.')

add_heading(doc, 'File-Based Query Cache', 2)
doc.add_paragraph(
    'In addition to SQLite, API search responses are cached to disk at '
    'agents/mabool/api/cache/. '
    'Cache key = SHA256 hash of (query, operation_mode, anchor_ids). '
    'TTL = 600 seconds. Format = JSON. Identical queries return instantly from cache.'
)

# ─── SECTION 9: CONFIGURATION ─────────────────────────────────────────────────
add_heading(doc, '9. Configuration Files and Environment Setup', 1)

add_heading(doc, 'Configuration Hierarchy', 2)
add_code_block(doc,
"""config.toml  (base defaults)
    ↓
config.extra.*.toml  (overlays, applied alphabetically)
    ↓
.env.secret  (secrets; overwrites matching keys)
    ↓
ConfigSettings  (in-memory at runtime)
    ↓
context-var  (per-request isolation via Python contextvars)""")

add_heading(doc, 'Key config.toml Settings', 2)
tbl7 = doc.add_table(rows=1, cols=3)
tbl7.style = 'Table Grid'
tbl7.rows[0].cells[0].text = 'Section'
tbl7.rows[0].cells[1].text = 'Key'
tbl7.rows[0].cells[2].text = 'Value / Purpose'
cfg_rows = [
    ('[default]', 'cache.enabled / ttl', 'true / 600 s — query result caching'),
    ('[default.relevance_judgement]', 'relevance_model_name', 'openai:gpt5mini-minimal-reasoning-default'),
    ('[default.relevance_judgement]', 'openai_concurrency', '75 — parallel scoring requests'),
    ('[default.relevance_judgement]', 'quota', '250 — max papers scored (cost control)'),
    ('[default.relevance_judgement]', 'initial_batch_size / growth_factor', '20 / 2 — dynamic batching'),
    ('[default.s2_api]', 'concurrency / timeout', '10 / 20 s — S2 API rate control'),
    ('[default.s2_api]', 'total_papers_limit', '1000 — max candidates per search'),
    ('[default.snowball_agent]', 'forward_top_k / backward_top_k', '200 / 200 — citation graph breadth'),
    ('[default.llm_abstraction]', 'gpt4o_default_model', 'gpt-4o-2024-11-20'),
    ('[default.llm_abstraction]', 'gemini2flash_default_model', 'gemini-2.0-flash'),
    ('[default.llm_abstraction]', 'temperature', '1.0 (deterministic)'),
]
for r in cfg_rows:
    add_table_row(tbl7, r)
doc.add_paragraph('')

add_heading(doc, 'Required API Keys (.env.secret)', 2)
tbl8 = doc.add_table(rows=1, cols=3)
tbl8.style = 'Table Grid'
tbl8.rows[0].cells[0].text = 'Variable'
tbl8.rows[0].cells[1].text = 'Required?'
tbl8.rows[0].cells[2].text = 'Purpose'
keys = [
    ('S2_API_KEY', 'REQUIRED (fatal)', 'Semantic Scholar paper search'),
    ('OPENAI_API_KEY', 'Required', 'Query analysis, relevance judgment, summaries'),
    ('GOOGLE_API_KEY', 'Required', 'Gemini-based specific paper matching'),
    ('COHERE_API_KEY', 'Optional', 'Cohere reranking (skipped if absent)'),
]
for r in keys:
    add_table_row(tbl8, r)
doc.add_paragraph('')

# ─── SECTION 10: HOW TO RUN ───────────────────────────────────────────────────
add_heading(doc, '10. How to Run / Build the Project', 1)

add_heading(doc, 'Prerequisites', 2)
doc.add_paragraph('1. Python 3.12.8+ or 3.13.x')
doc.add_paragraph('2. uv package manager: curl -LsSf https://astral.sh/uv/install.sh | sh')
doc.add_paragraph('3. API keys for Semantic Scholar, OpenAI, and Google Gemini (Cohere optional)')

add_heading(doc, 'Installation', 2)
add_code_block(doc,
"""# 1. Clone the repository
git clone https://github.com/allenai/asta-paper-finder

# 2. Sync all Python dependencies
make sync-dev

# 3. Create secrets file
cat > agents/mabool/api/conf/.env.secret << EOF
S2_API_KEY=your_s2_api_key
OPENAI_API_KEY=your_openai_key
GOOGLE_API_KEY=your_google_key
COHERE_API_KEY=your_cohere_key_optional
EOF""")

add_heading(doc, 'Running the Application', 2)
add_code_block(doc,
"""# Method 1: Recommended (via Makefile)
cd agents/mabool/api
make start-dev

# Method 2: Manual Gunicorn
APP_CONFIG_ENV=dev .venv/bin/gunicorn \\
    -k uvicorn.workers.UvicornWorker \\
    --workers 1 --timeout 0 --bind 0.0.0.0:8000 --reload \\
    'mabool.api.app:create_app()'

# Method 3: Uvicorn (debug mode)
APP_CONFIG_ENV=dev .venv/bin/uvicorn \\
    mabool.api.app:create_app --reload --host 0.0.0.0 --port 8000""")

add_heading(doc, 'Access Points (default: localhost:8000)', 2)
doc.add_paragraph('• Web UI:         http://localhost:8000/ui/')
doc.add_paragraph('• Swagger Docs:   http://localhost:8000/docs')
doc.add_paragraph('• Health Check:   http://localhost:8000/health')
doc.add_paragraph('• Search API:     POST http://localhost:8000/api/2/rounds')

add_heading(doc, 'Developer Workflow', 2)
add_code_block(doc,
"""make type-check    # Pyright static type checking
make lint          # Ruff + Flake8
make format        # Auto-format with Ruff
make fix           # Auto-fix lint issues
make test          # Run pytest suite
make test-cov      # With coverage report
make style         # format + lint + type-check combined""")

# ─── SECTION 11: KEY ALGORITHMS ───────────────────────────────────────────────
add_heading(doc, '11. Key Algorithms and Business Logic', 1)

add_heading(doc, 'Query Analysis Algorithm', 2)
doc.add_paragraph(
    'Input: Natural language string (e.g., "transformer papers from MIT before 2022").'
)
doc.add_paragraph(
    'Process: A structured LLM prompt extracts: (1) content keywords, (2) author names + '
    'affiliations, (3) venue names, (4) time range, (5) field of study, '
    '(6) recency/centrality preference, (7) broad vs. specific query type. '
    'Error recovery handles conflicting flags, non-actionable queries, and non-paper-finding intents (refusals).'
)

add_heading(doc, 'Relevance Judgment Algorithm', 2)
add_code_block(doc,
"""Dynamic Batching:
  Iteration 1: batch_size = 20  papers
  Iteration 2: batch_size = 20 × 2 = 40
  Iteration 3: batch_size = 40 × 2 = 80
  ... (capped at quota = 250)

Concurrency: 75 parallel LLM requests
LLM prompt: "Score this paper 0-1 for relevance to the query."
Output per paper: relevance_judgement.score ∈ [0, 1]""")
doc.add_paragraph(
    'Starting small allows early failure detection for invalid queries. '
    'The batch grows as confidence in valid results increases.'
)

add_heading(doc, 'Multi-Criteria Sorting Algorithm', 2)
add_code_block(doc,
"""final_agent_score =
    w_relevance  × normalize(relevance_judgement.score)
  + w_recency    × normalize(year)
  + w_centrality × normalize(log(citation_count + 1))
  + w_order      × (1 / original_rank)

Each dimension normalized to [0, 1].
Weights are configurable per query type and user preference.
Papers sorted descending by final_agent_score.""")

add_heading(doc, 'Broad Search Algorithm (Diligent Mode — up to 3 iterations)', 2)
add_code_block(doc,
"""Iteration 1:
  LLM generates N=5 initial keyword search queries
  → S2 API searches (parallel, concurrency=10)
  → Collect ~50-200 papers
  → Relevance judgment

Iteration 2 (if results insufficient):
  LLM analyzes prior results
  → Suggests N=5 refined queries
  → Merge with previous

Iteration 3: Similar refinement

Termination: high-relevance threshold reached OR max_iterations hit
Return: aggregated, deduplicated collection""")

add_heading(doc, 'Fast Search Algorithm (Parallel)', 2)
add_code_block(doc,
"""DenseAgent (Vespa vector search) ─────┐
SnowballAgent (citation graph)  ──────┤──► Merge & Deduplicate ──► Judgment ──► Sort
BroadSearchByKeywordAgent       ──────┘

All three execute concurrently (async).
Results merged via DocumentCollection.merged() (union, dedup by corpus_id).
Single judgment phase on the combined set.
Typical runtime: ~30 seconds.""")

add_heading(doc, 'Citation Snowball Algorithm', 2)
doc.add_paragraph(
    'Forward Citations: S2 API GET /paper/{corpus_id}/citations → up to 200 papers citing the anchor. '
    'Backward Citations: S2 API GET /paper/{corpus_id}/references → up to 200 papers cited by the anchor. '
    'Use case: starting from a known related paper, expand to its neighborhood in the citation graph.'
)

# ─── SECTION 12: DESIGN PATTERNS ─────────────────────────────────────────────
add_heading(doc, '12. Notable Design Patterns', 1)

patterns = [
    ('1. Operative Pattern (Custom Agent Framework)',
     'Operative[INPUT, OUTPUT, STATE] base class. Every agent is an async state machine that '
     'returns CompleteResponse, PartialResponse, or VoidResponse. Agents compose by creating '
     'child operatives. Enables structured error handling, scope cleanup, and type-safe I/O across '
     'the 14 search agents.'),
    ('2. Dependency Injection (Custom ai2i.di)',
     'Providers annotated with @module.provides(scope="singleton"|"round"). Dependencies declared '
     'via DI.requires() and DI.config(). The DI framework resolves the entire graph at startup '
     '(singletons) or at request time (round scope). Makes every component independently testable '
     'by substituting mock providers.'),
    ('3. LangChain Integration',
     'define_chat_llm_call() builds typed LangChain chains for OpenAI and Google Gemini. '
     'MaboolCallbackHandler intercepts on_llm_end events to accumulate token usage per model, '
     'enabling per-request cost attribution.'),
    ('4. Document Collection with Lazy Loading',
     'DocumentCollection fields are loaded on demand via with_fields(). Only corpus_id and title '
     'are fetched initially; abstract, authors, citations, etc. are fetched only when needed. '
     'This minimizes S2 API calls and speeds up early pipeline stages.'),
    ('5. Configuration Context Variables',
     'Python contextvars provide per-request config isolation. Each async request can override '
     'config values without affecting other concurrent requests. Essential for supporting '
     'different operation_modes (fast vs. diligent) in parallel.'),
    ('6. File-Based Query Cache',
     'Identical queries (same description + mode + anchors) return cached JSON within 600 seconds, '
     'skipping all LLM calls. Cache key = SHA256 of parameters. Dramatically reduces API costs '
     'for repeated or near-identical queries.'),
    ('7. Atomic SQLite Operations',
     'History inserts use INSERT ... SELECT ... WHERE EXISTS to atomically validate the parent '
     'session exists and insert in a single statement. Eliminates TOCTOU race conditions without '
     'explicit locking.'),
    ('8. Priority Semaphore',
     'A PrioritySemaphore(concurrency=3) gates all search executions globally. Prevents server '
     'overload and provides FIFO fairness when multiple users search simultaneously.'),
    ('9. Callback-Based Token Tracking',
     'LangChain AsyncCallbackHandler subclass collects usage_metadata from every LLM response. '
     'Token counts by model (prompt, completion, reasoning) are returned in every API response '
     'and aggregated in /api/stats for cost monitoring.'),
    ('10. Monorepo with uv Workspace',
     'Six internal Python packages (libs/*) are co-versioned and co-developed. uv workspace '
     'manages all packages and their inter-dependencies. A root Makefile delegates to '
     'per-package Makefiles via make_foreach.sh for parallel execution.'),
]

for title, desc in patterns:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    doc.add_paragraph(desc)

# ─── SECTION 13: SUMMARY ──────────────────────────────────────────────────────
add_heading(doc, '13. Summary', 1)

doc.add_paragraph(
    'ASTA Paper Finder is a sophisticated multi-tier system that combines modern Python '
    'async programming, multiple LLM providers, the Semantic Scholar academic graph, and '
    'a carefully designed agent architecture to deliver accurate, explainable paper recommendations.'
)

summary_points = [
    ('Frontend', 'Single-file HTML UI with session management, bookmarks, notes, and token stats.'),
    ('API Layer', 'FastAPI with Pydantic validation, priority semaphore, and file-based caching.'),
    ('Agent Framework', 'Custom Operative pattern: 14 specialized, composable search agents.'),
    ('LLM Integration', 'Unified abstraction (ai2i.chain) over OpenAI and Google Gemini; Cohere reranking.'),
    ('Data Management', 'Lazy-loaded DocumentCollections; Semantic Scholar as the authoritative data source.'),
    ('Infrastructure', 'Custom DI framework, TOML config, context-var isolation, token tracking.'),
    ('Persistence', 'SQLite for session history and bookmarks; file cache for query results.'),
    ('Deployment Target', 'Local development and research; single worker; no authentication layer.'),
]

tbl9 = doc.add_table(rows=1, cols=2)
tbl9.style = 'Table Grid'
tbl9.rows[0].cells[0].text = 'Component'
tbl9.rows[0].cells[1].text = 'Description'
for r in summary_points:
    add_table_row(tbl9, r)
doc.add_paragraph('')

# ─── SAVE ─────────────────────────────────────────────────────────────────────
output_path = '/Users/ipanda/Documents/Code/asta/ASTA_Technical_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
